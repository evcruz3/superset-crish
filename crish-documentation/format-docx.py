#!/usr/bin/env python3
"""
Post-processing script to format DOCX paragraphs with justification and spacing.
This script modifies the generated DOCX file to apply proper paragraph formatting.
"""

import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

def format_docx(filename):
    """Apply paragraph formatting to a DOCX file."""
    try:
        # Open the document
        doc = Document(filename)
        
        print(f"Processing {filename}...")
        print(f"Found {len(doc.paragraphs)} paragraphs")
        
        # Format all paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            # Skip empty paragraphs
            if not paragraph.text.strip():
                continue
                
            # Apply justification
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Set paragraph spacing
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(12)  # 12pt space after
            paragraph_format.line_spacing = 1.15   # 1.15 line spacing
            paragraph_format.first_line_indent = Pt(0)  # No first line indent
            
            # Set font if needed
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12)
        
        # Handle tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            paragraph.paragraph_format.space_after = Pt(6)
        
        # Handle image captions (center alignment)
        for paragraph in doc.paragraphs:
            # Check if paragraph contains figure caption text
            if any(keyword in paragraph.text.lower() for keyword in ['figure', 'fig.', 'image']):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(12)
                paragraph.paragraph_format.space_before = Pt(6)
                
                # Make caption slightly smaller and italic
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.italic = True
        
        # Save the formatted document
        doc.save(filename)
        print(f"✓ Successfully formatted {filename}")
        print("Applied:")
        print("  - Justified paragraph alignment")
        print("  - 12pt spacing after paragraphs")
        print("  - 1.15 line spacing")
        print("  - Centered figure captions")
        print("  - Arial 12pt font")
        
    except Exception as e:
        print(f"Error processing {filename}: {e}")
        sys.exit(1)

def main():
    if len(sys.argv) != 2:
        print("Usage: python format-docx.py <filename.docx>")
        sys.exit(1)
    
    filename = sys.argv[1]
    format_docx(filename)

if __name__ == "__main__":
    main()